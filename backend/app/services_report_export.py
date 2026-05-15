from __future__ import annotations
from pathlib import Path
from datetime import datetime
from math import pi

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt


def _style_run(run, bold: bool = False, size: int = 13):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold


def _style_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size: int = 13):
    paragraph.alignment = align
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    for run in paragraph.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)


def _add_kv_paragraph(doc: Document, label: str, value: str):
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    _style_run(run, bold=True)
    value_run = p.add_run(str(value or ''))
    _style_run(value_run)
    _style_paragraph(p)


def _format_document(doc: Document):
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1)
    for p in doc.paragraphs:
        _style_paragraph(p)
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _style_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER)


def _set_heading(paragraph, size: int = 13, all_caps: bool = False):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if all_caps else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    for run in paragraph.runs:
        text = run.text.upper() if all_caps else run.text
        run.text = text
        _style_run(run, bold=True, size=size)


def _chart_paths(base_dir: Path, safe_name: str, stamp: str, prefix: str = 'report') -> tuple[Path, Path]:
    chart_dir = base_dir / 'charts'
    chart_dir.mkdir(parents=True, exist_ok=True)
    bar_path = chart_dir / f"{prefix}_bar_{safe_name}_{stamp}.png"
    radar_path = chart_dir / f"{prefix}_radar_{safe_name}_{stamp}.png"
    return bar_path, radar_path


def _build_bar_chart(inputs: dict, output_path: Path):
    labels = ['X1', 'X2', 'X3', 'X4', 'X5', 'X6']
    values = [float(inputs.get(k, 0) or 0) for k in labels]
    fig, ax = plt.subplots(figsize=(8, 4.8))
    bars = ax.bar(labels, values, color=['#2563eb', '#1d4ed8', '#3b82f6', '#60a5fa', '#0ea5e9', '#0284c7'])
    ax.set_ylim(0, 5)
    ax.set_title('Biểu đồ cột các chỉ tiêu tổng hợp nội bộ', fontsize=12, fontweight='bold')
    ax.set_ylabel('Giá trị chỉ tiêu tổng hợp')
    ax.grid(axis='y', linestyle='--', alpha=0.35)
    for bar, value in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width() / 2, value + 0.05, f'{value:.2f}', ha='center', va='bottom', fontsize=9)
    fig.tight_layout()
    fig.savefig(output_path, dpi=180)
    plt.close(fig)


def _build_radar_chart(inputs: dict, output_path: Path):
    labels = ['X1', 'X2', 'X3', 'X4', 'X5', 'X6']
    values = [float(inputs.get(k, 0) or 0) for k in labels]
    values += values[:1]
    angles = [n / float(len(labels)) * 2 * pi for n in range(len(labels))]
    angles += angles[:1]
    fig = plt.figure(figsize=(6, 6))
    ax = plt.subplot(111, polar=True)
    ax.set_theta_offset(pi / 2)
    ax.set_theta_direction(-1)
    plt.xticks(angles[:-1], labels)
    ax.set_rlabel_position(0)
    plt.yticks([1, 2, 3, 4, 5], ['1', '2', '3', '4', '5'], color='gray', size=8)
    plt.ylim(0, 5)
    ax.plot(angles, values, linewidth=2, linestyle='solid', color='#2563eb')
    ax.fill(angles, values, color='#60a5fa', alpha=0.35)
    plt.title('Biểu đồ radar các chỉ tiêu tổng hợp nội bộ', y=1.08, fontsize=12, fontweight='bold')
    fig.tight_layout()
    fig.savefig(output_path, dpi=180)
    plt.close(fig)


def _build_training_loss_chart(history: list, output_path: Path):
    epochs = [item.get('epoch') for item in history]
    train_loss = [item.get('train_loss', 0) for item in history]
    test_loss = [item.get('test_loss', 0) for item in history]
    fig, ax = plt.subplots(figsize=(8, 4.8))
    ax.plot(epochs, train_loss, marker='o', label='Train loss', color='#2563eb')
    ax.plot(epochs, test_loss, marker='s', label='Test loss', color='#d97706')
    ax.set_title('Diễn biến loss theo epoch', fontsize=12, fontweight='bold')
    ax.set_xlabel('Epoch')
    ax.set_ylabel('Loss')
    ax.grid(True, linestyle='--', alpha=0.35)
    ax.legend()
    fig.tight_layout()
    fig.savefig(output_path, dpi=180)
    plt.close(fig)


def build_report_docx(payload: dict, report_dir: str) -> tuple[str, str]:
    out_dir = Path(report_dir).resolve()
    out_dir.mkdir(parents=True, exist_ok=True)
    safe_project = str(payload.get('projectName') or 'coastalrisk-report').replace(' ', '_')
    stamp = datetime.utcnow().strftime('%Y%m%d%H%M%S')
    file_name = f"report_{safe_project}_{stamp}.docx"
    output_path = out_dir / file_name
    bar_chart_path, radar_chart_path = _chart_paths(out_dir, safe_project, stamp)
    doc = Document()
    title = doc.add_heading('BÁO CÁO DỰ BÁO RỦI RO CÔNG TRÌNH ĐÊ, KÈ BIỂN', level=0)
    _set_heading(title, size=14, all_caps=True)
    subtitle = doc.add_paragraph('Hệ thống CoastalRisk WebApp v1 - ANFIS / Sugeno / TensorFlow')
    _style_paragraph(subtitle, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_heading('1. Thông tin chung', level=1)
    _add_kv_paragraph(doc, 'Ngày báo cáo', payload.get('reportDate'))
    _add_kv_paragraph(doc, 'Công trình', payload.get('projectName'))
    _add_kv_paragraph(doc, 'Địa điểm', payload.get('projectLocation'))
    _add_kv_paragraph(doc, 'Vùng nghiên cứu', payload.get('region'))
    _add_kv_paragraph(doc, 'Giai đoạn thi công', payload.get('projectStage'))
    _add_kv_paragraph(doc, 'Mô hình sử dụng', payload.get('modelName'))
    _add_kv_paragraph(doc, 'Loại mô hình', payload.get('modelType'))
    doc.add_heading('2. Kết quả dự báo', level=1)
    result = payload.get('result') or {}
    _add_kv_paragraph(doc, 'Điểm rủi ro', result.get('score'))
    _add_kv_paragraph(doc, 'Mức rủi ro', result.get('level'))
    risk_matrix = payload.get('riskMatrix') or {}
    _add_kv_paragraph(doc, 'Likelihood', risk_matrix.get('likelihood'))
    _add_kv_paragraph(doc, 'Impact', risk_matrix.get('impact'))
    _add_kv_paragraph(doc, 'Matrix Score', risk_matrix.get('matrixScore'))
    doc.add_heading('3. Bảng chỉ tiêu tổng hợp nội bộ phục vụ dự báo', level=1)
    inputs = payload.get('inputs') or {}
    input_table = doc.add_table(rows=1, cols=3)
    input_table.style = 'Table Grid'
    hdr = input_table.rows[0].cells
    hdr[0].text = 'Biến'; hdr[1].text = 'Giá trị'; hdr[2].text = 'Diễn giải'
    labels = {'X1': 'Điều kiện tự nhiên', 'X2': 'Địa chất - nền móng', 'X3': 'Kỹ thuật - công nghệ', 'X4': 'Vật liệu - logistics', 'X5': 'Tổ chức - quản lý', 'X6': 'An toàn - môi trường'}
    for key in ['X1', 'X2', 'X3', 'X4', 'X5', 'X6']:
        row = input_table.add_row().cells
        row[0].text = key; row[1].text = str(inputs.get(key, '')); row[2].text = labels.get(key, '')
    doc.add_heading('4. Chỉ số đánh giá mô hình', level=1)
    metrics = payload.get('metrics') or {}
    metric_table = doc.add_table(rows=1, cols=3)
    metric_table.style = 'Table Grid'
    hdr = metric_table.rows[0].cells
    hdr[0].text = 'Thông số'; hdr[1].text = 'Giá trị'; hdr[2].text = 'Nhận xét'
    for key, label, comment in [('mae', 'MAE', 'Sai số tuyệt đối trung bình'), ('rmse', 'RMSE', 'Độ lệch chuẩn sai số'), ('r2', 'R²', 'Khả năng giải thích biến thiên')]:
        row = metric_table.add_row().cells
        row[0].text = label; row[1].text = str(metrics.get(key, '')); row[2].text = comment
    doc.add_heading('5. Phân tích nhóm rủi ro nổi trội', level=1)
    dominant_factor = payload.get('dominantFactor')
    dominant_labels = {'X1': 'Nhóm 1 - Điều kiện tự nhiên', 'X2': 'Nhóm 2 - Địa chất - nền móng', 'X3': 'Nhóm 3 - Kỹ thuật - công nghệ', 'X4': 'Nhóm 4 - Vật liệu - logistics', 'X5': 'Nhóm 5 - Tổ chức - quản lý', 'X6': 'Nhóm 6 - An toàn - môi trường'}
    if dominant_factor:
        doc.add_paragraph(f"Nhóm rủi ro nổi trội nhất của công trình là {dominant_labels.get(dominant_factor, dominant_factor)} ({dominant_factor}).")
    doc.add_heading('6. Tóm tắt và khuyến nghị quản lý', level=1)
    doc.add_paragraph(str(payload.get('summary') or ''))
    recommendations = payload.get('recommendations') or {}
    immediate_actions = recommendations.get('immediateActions') or []
    priority_actions = recommendations.get('priorityActions') or []
    monitoring_actions = recommendations.get('monitoringActions') or []
    stakeholder_actions = recommendations.get('stakeholderActions') or {}
    management_message = recommendations.get('managementMessage')
    if management_message:
        doc.add_paragraph(str(management_message))
    if immediate_actions:
        doc.add_paragraph('Hành động cần ưu tiên ngay:')
        for item in immediate_actions:
            doc.add_paragraph(str(item), style='List Bullet')
    if priority_actions:
        doc.add_paragraph('Biện pháp kiểm soát trọng tâm:')
        for item in priority_actions:
            doc.add_paragraph(str(item), style='List Bullet')
    if monitoring_actions:
        doc.add_paragraph('Nội dung cần theo dõi liên tục:')
        for item in monitoring_actions:
            doc.add_paragraph(str(item), style='List Bullet')
    if stakeholder_actions:
        doc.add_paragraph('Kiến nghị theo nhóm đối tượng sử dụng:')
        if stakeholder_actions.get('siteCommand'):
            doc.add_paragraph('Ban chỉ huy công trường:')
            for item in stakeholder_actions.get('siteCommand', []):
                doc.add_paragraph(str(item), style='List Bullet')
        if stakeholder_actions.get('supervisionConsultant'):
            doc.add_paragraph('Tư vấn giám sát:')
            for item in stakeholder_actions.get('supervisionConsultant', []):
                doc.add_paragraph(str(item), style='List Bullet')
        if stakeholder_actions.get('investorPMU'):
            doc.add_paragraph('Chủ đầu tư / Ban QLDA:')
            for item in stakeholder_actions.get('investorPMU', []):
                doc.add_paragraph(str(item), style='List Bullet')
    doc.add_heading('7. Biểu đồ trực quan phục vụ phân tích rủi ro', level=1)
    _build_bar_chart(inputs, bar_chart_path)
    _build_radar_chart(inputs, radar_chart_path)
    p1 = doc.add_paragraph('Hình 1. Biểu đồ cột các chỉ tiêu tổng hợp nội bộ của hồ sơ công trình')
    _style_paragraph(p1, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_picture(str(bar_chart_path), width=Inches(6.2))
    p2 = doc.add_paragraph('Hình 2. Biểu đồ radar các chỉ tiêu tổng hợp nội bộ của hồ sơ công trình')
    _style_paragraph(p2, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_picture(str(radar_chart_path), width=Inches(5.8))
    doc.add_heading('8. Nhận xét học thuật', level=1)
    doc.add_paragraph('Báo cáo này được sinh tự động từ hệ thống CoastalRisk ANFIS nhằm hỗ trợ phân tích định lượng rủi ro thi công công trình đê, kè biển trên cơ sở các chỉ tiêu tổng hợp nội bộ và kết quả suy diễn của mô hình.')
    _format_document(doc)
    doc.save(output_path)
    return file_name, str(output_path)


def build_membership_compare_report_docx(payload: dict, report_dir: str) -> tuple[str, str]:
    out_dir = Path(report_dir).resolve()
    out_dir.mkdir(parents=True, exist_ok=True)
    stamp = datetime.utcnow().strftime('%Y%m%d%H%M%S')
    file_name = f"membership_compare_{stamp}.docx"
    output_path = out_dir / file_name
    doc = Document()
    title = doc.add_heading('BÁO CÁO SO SÁNH CÁC HÀM THÀNH VIÊN ANFIS', level=0)
    _set_heading(title, size=14, all_caps=True)
    _add_kv_paragraph(doc, 'Ngày báo cáo', payload.get('reportDate'))
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Membership'; hdr[1].text = 'Số model'; hdr[2].text = 'Model tốt nhất'; hdr[3].text = 'R²'; hdr[4].text = 'MAE / RMSE'
    for item in payload.get('groups', []):
        row = table.add_row().cells
        row[0].text = str(item.get('membership'))
        row[1].text = str(item.get('count'))
        row[2].text = str(item.get('bestModel'))
        row[3].text = str(item.get('r2'))
        row[4].text = f"{item.get('mae')} / {item.get('rmse')}"
    _format_document(doc)
    doc.save(output_path)
    return file_name, str(output_path)


def build_training_report_docx(payload: dict, report_dir: str) -> tuple[str, str]:
    out_dir = Path(report_dir).resolve()
    out_dir.mkdir(parents=True, exist_ok=True)
    safe_name = str(payload.get('modelName') or 'training-report').replace(' ', '_')
    stamp = datetime.utcnow().strftime('%Y%m%d%H%M%S')
    file_name = f"training_{safe_name}_{stamp}.docx"
    output_path = out_dir / file_name
    loss_chart_path, _ = _chart_paths(out_dir, safe_name, stamp, prefix='training')
    metrics = payload.get('metrics') or {}
    artifact = payload.get('artifact') or {}
    history = metrics.get('history') or artifact.get('history') or []
    _build_training_loss_chart(history, loss_chart_path)
    doc = Document()
    title = doc.add_heading('BÁO CÁO HUẤN LUYỆN MÔ HÌNH ANFIS', level=0)
    _set_heading(title, size=14, all_caps=True)
    _add_kv_paragraph(doc, 'Ngày báo cáo', payload.get('reportDate'))
    _add_kv_paragraph(doc, 'Tên mô hình', payload.get('modelName'))
    _add_kv_paragraph(doc, 'Vùng', payload.get('region'))
    _add_kv_paragraph(doc, 'Số epoch thực tế', metrics.get('epochs'))
    _add_kv_paragraph(doc, 'Kiểu khởi tạo luật', payload.get('ruleInitMode') or artifact.get('rule_init') or metrics.get('rule_init_mode'))
    _add_kv_paragraph(doc, 'Best epoch', metrics.get('best_epoch'))
    _add_kv_paragraph(doc, 'Best test loss', metrics.get('best_test_loss'))
    _add_kv_paragraph(doc, 'Early stopping', 'Có' if metrics.get('early_stopped') else 'Không')
    doc.add_heading('1. Chỉ số mô hình', level=1)
    _add_kv_paragraph(doc, 'MAE', metrics.get('mae'))
    _add_kv_paragraph(doc, 'RMSE', metrics.get('rmse'))
    _add_kv_paragraph(doc, 'R²', metrics.get('r2'))
    doc.add_heading('2. Diễn biến quá trình hội tụ theo epoch', level=1)
    p = doc.add_paragraph('Hình 1. Diễn biến loss theo epoch trong quá trình huấn luyện mô hình')
    _style_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_picture(str(loss_chart_path), width=Inches(6.2))
    doc.add_heading('3. Bảng luật tại mô hình tốt nhất', level=1)
    rules = artifact.get('rules') or []
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Luật'; hdr[1].text = 'Điều kiện'; hdr[2].text = 'Consequent'; hdr[3].text = 'Trạng thái'
    for rule in rules:
        row = table.add_row().cells
        row[0].text = str(rule.get('name'))
        row[1].text = ', '.join([f"{k} {v[0]} (c={v[1]}, σ={v[2]})" for k, v in (rule.get('inputs') or {}).items()])
        row[2].text = ' + '.join([f"{v}·{k}" for k, v in (rule.get('coeffs') or {}).items()]) + f" + {rule.get('bias', 0)}"
        origin = rule.get('origin') or artifact.get('rule_init') or payload.get('ruleInitMode') or ''
        row[3].text = f"{rule.get('status', 'Learned')} | {origin}".strip(' |')
    _format_document(doc)
    doc.save(output_path)
    return file_name, str(output_path)
